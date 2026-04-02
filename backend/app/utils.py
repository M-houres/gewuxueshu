import hashlib
import random
import re
import string
from pathlib import Path

from docx import Document
from pypdf import PdfReader


def make_invite_code(user_id: int) -> str:
    digest = hashlib.sha256(f"user-{user_id}".encode("utf-8")).hexdigest().upper()
    return digest[:8]


def make_order_no() -> str:
    prefix = "OD"
    rand = "".join(random.choices(string.digits, k=12))
    return f"{prefix}{rand}"


def gen_code() -> str:
    return "".join(random.choices(string.digits, k=6))


def is_phone_valid(phone: str) -> bool:
    return bool(re.fullmatch(r"1\d{10}", phone))


def safe_filename(name: str) -> str:
    allow = set(string.ascii_letters + string.digits + "._-")
    return "".join(ch if ch in allow else "_" for ch in name)


def detect_file_magic(path: Path) -> str:
    with path.open("rb") as f:
        head = f.read(8)
    if head.startswith(b"%PDF"):
        return ".pdf"
    if head.startswith(b"PK\x03\x04"):
        return ".docx"
    try:
        content = path.read_text(encoding="utf-8")
        if content:
            return ".txt"
    except Exception:
        pass
    return ""


def extract_text_from_file(path: Path) -> str:
    suffix = path.suffix.lower()
    if suffix == ".txt":
        return path.read_text(encoding="utf-8", errors="ignore")
    if suffix == ".docx":
        doc = Document(str(path))
        parts: list[str] = [p.text for p in doc.paragraphs if p.text]
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text = cell.text.strip()
                    if text:
                        parts.append(text)
        return "\n".join(parts)
    if suffix == ".pdf":
        reader = PdfReader(str(path))
        parts: list[str] = []
        for page in reader.pages:
            parts.append(page.extract_text() or "")
        return "\n".join(parts)
    raise ValueError("unsupported file type")


def count_billable_chars(text: str) -> int:
    filtered = [ch for ch in text if ch.isalnum() or ("\u4e00" <= ch <= "\u9fff")]
    return len(filtered)
