import re
from dataclasses import dataclass
from datetime import datetime
from difflib import SequenceMatcher
from pathlib import Path

from docx import Document
from sqlalchemy.orm import Session

from app.config import get_settings
from app.models import LLMErrorLog, SwitchLog, SystemSwitch, TaskType
from app.services.algo_package_service import run_active_package
from app.services.llm_service import generate_with_llm, load_llm_config
from app.utils import count_billable_chars, extract_text_from_file

MODE_LLM_PLUS_ALGO = "LLM_PLUS_ALGO"
MODE_ALGO_ONLY = "ALGO_ONLY"
settings = get_settings()


@dataclass
class ProcessResult:
    output_path: str
    result_json: dict


class ProcessingEngine:
    def __init__(self, db: Session) -> None:
        self.db = db
        self._current_switch: SystemSwitch | None = None
        self._current_task_id: int | None = None
        self._pipeline_usage = {"llm_used": False, "algo_package_used": False}

    def _get_or_init_switch(self) -> SystemSwitch:
        switch = self.db.query(SystemSwitch).first()
        if switch:
            return switch
        switch = SystemSwitch(
            current_mode=MODE_LLM_PLUS_ALGO if settings.llm_enabled_default else MODE_ALGO_ONLY,
            llm_enabled=settings.llm_enabled_default,
            llm_fail_count=0,
            llm_fail_threshold=3,
        )
        self.db.add(switch)
        self.db.flush()
        return switch

    def _switch_mode(self, target_mode: str, reason: str) -> None:
        switch = self._get_or_init_switch()
        if switch.current_mode == target_mode:
            return
        self.db.add(SwitchLog(from_mode=switch.current_mode, to_mode=target_mode, reason=reason))
        switch.current_mode = target_mode
        self.db.flush()

    def process(
        self,
        task_type: TaskType,
        platform: str,
        input_path: Path,
        output_path: Path,
        task_id: int | None = None,
        report_path: Path | None = None,
    ) -> ProcessResult:
        switch = self._get_or_init_switch()
        self._current_switch = switch
        self._current_task_id = task_id
        self._pipeline_usage = {"llm_used": False, "algo_package_used": False}

        llm_cfg = load_llm_config(self.db)
        switch.llm_enabled = bool(llm_cfg.get("enabled", False))
        normalized_platform = (platform or "").strip().lower()
        if not switch.llm_enabled:
            self._switch_mode(MODE_ALGO_ONLY, "llm disabled")
        elif switch.current_mode != MODE_ALGO_ONLY or switch.llm_fail_count < switch.llm_fail_threshold:
            self._switch_mode(MODE_LLM_PLUS_ALGO, "llm healthy")

        source_text = extract_text_from_file(input_path)
        report_text = self._load_optional_report(report_path)
        report_summary = self._extract_report_summary(task_type, report_text)

        if task_type == TaskType.AIGC_DETECT:
            algo_result = self._run_algo_package(normalized_platform, task_type, source_text)
            detect_result = self._build_detect_result(
                text=source_text,
                platform=normalized_platform,
                mode=switch.current_mode,
                report_summary=report_summary,
                algo_result=algo_result,
            )
            self._write_detect_report_pdf(output_path, detect_result)
            return ProcessResult(output_path=str(output_path), result_json=detect_result)

        if input_path.suffix.lower() == ".docx":
            self._transform_docx(input_path, output_path, task_type, normalized_platform, report_summary)
            output_text = extract_text_from_file(output_path)
        else:
            output_text = self._transform_text(source_text, task_type, normalized_platform, report_summary)
            output_path.write_text(output_text, encoding="utf-8")

        result_json = self._build_transform_result(
            task_type=task_type,
            platform=normalized_platform,
            mode=switch.current_mode,
            source_text=source_text,
            output_text=output_text,
            report_summary=report_summary,
        )
        return ProcessResult(output_path=str(output_path), result_json=result_json)

    def _load_optional_report(self, report_path: Path | None) -> str:
        if report_path is None or not report_path.exists():
            return ""
        try:
            return extract_text_from_file(report_path)
        except Exception:
            return ""

    def _heuristic_ai_score(self, text: str) -> float:
        clean = text.strip()
        if not clean:
            return 0.0
        normalized = clean.replace("。", ".").replace("！", ".").replace("？", ".")
        sentences = [seg.strip() for seg in normalized.split(".") if seg.strip()]
        if not sentences:
            return 0.0
        avg_len = sum(len(s) for s in sentences) / len(sentences)
        uniq_ratio = len(set(clean)) / max(len(clean), 1)
        score = min(1.0, max(0.0, (avg_len / 80.0) * 0.6 + (1 - uniq_ratio) * 0.4))
        return round(score, 4)

    def _iter_body_runs(self, doc: Document):
        for para in doc.paragraphs:
            for run in para.runs:
                yield run
        if settings.docx_process_table_text:
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for para in cell.paragraphs:
                            for run in para.runs:
                                yield run

    def _transform_docx(
        self,
        input_path: Path,
        output_path: Path,
        task_type: TaskType,
        platform: str,
        report_summary: dict | None = None,
    ) -> None:
        doc = Document(str(input_path))
        summary = report_summary or {}
        for run in self._iter_body_runs(doc):
            if run.text:
                run.text = self._transform_text(run.text, task_type, platform, summary)
        doc.save(str(output_path))

    def _run_algo_package(self, platform: str, task_type: TaskType, text: str):
        try:
            result = run_active_package(
                self.db,
                platform=platform,
                function_type=task_type.value,
                text=text,
            )
        except Exception:
            return None
        if result is None:
            return None
        self._pipeline_usage["algo_package_used"] = True
        value, _meta = result
        return value

    def _run_llm(self, task_type: TaskType, text: str) -> str | None:
        switch = self._current_switch or self._get_or_init_switch()
        if switch.current_mode != MODE_LLM_PLUS_ALGO or not switch.llm_enabled:
            return None
        if task_type not in {TaskType.DEDUP, TaskType.REWRITE}:
            return None
        try:
            output = generate_with_llm(self.db, task_type=task_type, text=text)
            switch.llm_fail_count = 0
            self._pipeline_usage["llm_used"] = True
            self.db.flush()
            return output
        except Exception as exc:
            switch.llm_fail_count += 1
            should_downgrade = switch.llm_fail_count >= switch.llm_fail_threshold
            if should_downgrade:
                self._switch_mode(MODE_ALGO_ONLY, f"llm_fail:{str(exc)[:160]}")
            self.db.add(
                LLMErrorLog(
                    task_id=self._current_task_id,
                    error_type=exc.__class__.__name__,
                    error_detail=str(exc)[:500],
                    trigger_downgrade=should_downgrade,
                )
            )
            self.db.flush()
            return None

    def _transform_text(self, text: str, task_type: TaskType, platform: str, report_summary: dict) -> str:
        llm_output = self._run_llm(task_type, text)
        if isinstance(llm_output, str) and llm_output.strip():
            return llm_output

        algo_result = self._run_algo_package(platform, task_type, text)
        if isinstance(algo_result, str) and algo_result.strip():
            return algo_result
        if isinstance(algo_result, dict):
            candidate_keys = ("text", "rewritten_text", "output_text", "result_text")
            for key in candidate_keys:
                output = algo_result.get(key)
                if isinstance(output, str) and output.strip():
                    return output

        pressure = report_summary.get("pressure", "low")
        normalized = self._normalize_text(text)
        if task_type == TaskType.DEDUP:
            replacements = {
                "因此": "由此可见",
                "但是": "然而",
                "首先": "第一",
                "其次": "第二",
                "总之": "综上所述",
                "可以看出": "据此可见",
                "本文认为": "本文进一步指出",
            }
            output = self._apply_replacements(normalized, replacements)
            output = self._split_long_sentences(output, 48 if pressure == "high" else 64)
            return output
        if task_type == TaskType.REWRITE:
            replacements = {
                "研究表明": "已有研究指出",
                "可以看出": "据此可见",
                "非常": "较为",
                "重要": "关键",
                "很多": "大量",
                "我们发现": "研究发现",
                "这个": "该",
            }
            output = self._apply_replacements(normalized, replacements)
            output = self._split_long_sentences(output, 54 if pressure == "high" else 72)
            return output
        return normalized

    def _normalize_text(self, text: str) -> str:
        output = re.sub(r"[ \t]+", " ", text)
        output = re.sub(r"\n{3,}", "\n\n", output)
        return output.strip()

    def _apply_replacements(self, text: str, replacements: dict[str, str]) -> str:
        output = text
        for src, target in replacements.items():
            output = output.replace(src, target)
        return output

    def _split_long_sentences(self, text: str, threshold: int) -> str:
        chunks = re.split(r"([。！？!?])", text)
        rebuilt: list[str] = []
        for index in range(0, len(chunks), 2):
            sentence = chunks[index].strip()
            punct = chunks[index + 1] if index + 1 < len(chunks) else ""
            if len(sentence) > threshold and "，" in sentence:
                parts = [part.strip() for part in sentence.split("，") if part.strip()]
                current: list[str] = []
                current_len = 0
                groups: list[str] = []
                for part in parts:
                    if current and current_len + len(part) > threshold:
                        groups.append("，".join(current))
                        current = [part]
                        current_len = len(part)
                    else:
                        current.append(part)
                        current_len += len(part)
                if current:
                    groups.append("，".join(current))
                rebuilt.append("。".join(groups) + punct)
            else:
                rebuilt.append(sentence + punct)
        return "".join(rebuilt).strip()

    def _text_stats(self, text: str) -> dict:
        clean = text.strip()
        sentences = [part.strip() for part in re.split(r"[。！？!?；;\n]+", clean) if part.strip()]
        paragraphs = [part.strip() for part in clean.splitlines() if part.strip()]
        sentence_count = len(sentences)
        avg_sentence_length = round(sum(len(item) for item in sentences) / sentence_count, 2) if sentence_count else 0
        return {
            "char_count": count_billable_chars(clean),
            "paragraph_count": len(paragraphs),
            "sentence_count": sentence_count,
            "avg_sentence_length": avg_sentence_length,
        }

    def _extract_report_summary(self, task_type: TaskType, report_text: str) -> dict:
        content = " ".join((report_text or "").split())
        summary = {
            "available": bool(content),
            "metrics": [],
            "highlights": [],
            "recommended_actions": [],
            "pressure": "low",
        }
        if not content:
            summary["recommended_actions"] = ["未上传辅助报告，本次按正文通用策略处理。"]
            return summary

        if task_type == TaskType.DEDUP:
            total_ratio = self._extract_percent(content, ["总文字复制比", "全文总重复率", "重复率", "总复制比"])
            quote_ratio = self._extract_percent(content, ["去除引用复制比"])
            self_ratio = self._extract_percent(content, ["去除本人已发表文献复制比"])
            for label, value in (
                ("总文字复制比", total_ratio),
                ("去除引用复制比", quote_ratio),
                ("去除本人已发表文献复制比", self_ratio),
            ):
                if value is not None:
                    summary["metrics"].append({"label": label, "value": value, "unit": "%"})
            summary["highlights"] = [word for word in ["全文", "检测报告", "总文字复制比", "去除引用复制比"] if word in content][:4]
            if total_ratio is not None and total_ratio >= 25:
                summary["recommended_actions"].append("重复率偏高，优先处理定义、综述和结论性长句。")
                summary["pressure"] = "high"
            elif total_ratio is not None and total_ratio >= 15:
                summary["recommended_actions"].append("重复率中等，优先改写高频连接词和段首句。")
                summary["pressure"] = "medium"
            if quote_ratio is not None and quote_ratio >= 10:
                summary["recommended_actions"].append("检查引用说明是否过少，必要时补充规范引文表达。")
            if self_ratio is not None and self_ratio >= 10:
                summary["recommended_actions"].append("留意与本人历史文本重合的定义和结论段。")
        else:
            ai_ratio = self._extract_percent(content, ["AIGC总体风险", "总体风险", "AIGC疑似度", "AI生成疑似度", "疑似AI生成"])
            high_ratio = self._extract_percent(content, ["高风险占比", "高风险段落占比"])
            for label, value in (
                ("总体风险", ai_ratio),
                ("高风险占比", high_ratio),
            ):
                if value is not None:
                    summary["metrics"].append({"label": label, "value": value, "unit": "%"})
            summary["highlights"] = [word for word in ["AIGC", "疑似AI", "高风险段落", "全文"] if word.lower() in content.lower()][:4]
            if ai_ratio is not None and ai_ratio >= 50:
                summary["recommended_actions"].append("AIGC 风险偏高，优先拆分长句并弱化模板化连接词。")
                summary["pressure"] = "high"
            elif ai_ratio is not None and ai_ratio >= 30:
                summary["recommended_actions"].append("AIGC 风险中等，建议提升句式变化和论证层次。")
                summary["pressure"] = "medium"
            if high_ratio is not None and high_ratio >= 20:
                summary["recommended_actions"].append("重点复核高风险段落，尤其是定义句和总结句。")

        if not summary["recommended_actions"]:
            if task_type == TaskType.DEDUP:
                summary["recommended_actions"].append("建议重点复核连续长句、定义表述和文献综述段落。")
            else:
                summary["recommended_actions"].append("建议优先调整摘要、结论和高频模板化表达。")
        return summary

    def _extract_percent(self, text: str, keywords: list[str]) -> float | None:
        for keyword in keywords:
            pattern = rf"{re.escape(keyword)}[^0-9]{{0,12}}(\d+(?:\.\d+)?)\s*%"
            match = re.search(pattern, text, flags=re.IGNORECASE)
            if match:
                return round(float(match.group(1)), 2)
        return None


    def _clamp_score(self, value: float) -> float:
        return max(0.0, min(1.0, float(value)))

    def _platform_detect_profile(self, platform: str) -> dict:
        key = (platform or "").strip().lower()
        profiles = {
            "cnki": {
                "name": "cnki_like",
                "baseline_weight": 0.7,
                "style_weight": 0.2,
                "repeat_weight": 0.1,
                "offset": 0.0,
                "high": 0.65,
                "medium": 0.35,
            },
            "vip": {
                "name": "vip_like",
                "baseline_weight": 0.64,
                "style_weight": 0.24,
                "repeat_weight": 0.12,
                "offset": -0.02,
                "high": 0.62,
                "medium": 0.33,
            },
            "paperpass": {
                "name": "paperpass_like",
                "baseline_weight": 0.66,
                "style_weight": 0.16,
                "repeat_weight": 0.18,
                "offset": 0.03,
                "high": 0.6,
                "medium": 0.32,
            },
        }
        return profiles.get(key, profiles["cnki"])

    def _simulate_platform_detect_score(self, platform: str, text: str, base_score: float) -> tuple[float, dict, dict]:
        profile = self._platform_detect_profile(platform)
        stats = self._text_stats(text)
        clean = (text or "").strip()
        compact = " ".join(clean.split())
        unique_ratio = (len(set(compact)) / len(compact)) if compact else 1.0
        repeat_signal = self._clamp_score(1.0 - unique_ratio)
        avg_len = float(stats.get("avg_sentence_length") or 0.0)
        style_signal = self._clamp_score((avg_len - 18.0) / 55.0)

        weighted = (
            float(base_score) * profile["baseline_weight"]
            + style_signal * profile["style_weight"]
            + repeat_signal * profile["repeat_weight"]
            + profile["offset"]
        )
        score = round(self._clamp_score(weighted), 4)
        breakdown = {
            "base_score": round(float(base_score), 4),
            "style_signal": round(style_signal, 4),
            "repeat_signal": round(repeat_signal, 4),
            "weights": {
                "baseline": profile["baseline_weight"],
                "style": profile["style_weight"],
                "repeat": profile["repeat_weight"],
                "offset": profile["offset"],
            },
            "thresholds": {
                "high": profile["high"],
                "medium": profile["medium"],
            },
        }
        return score, profile, breakdown

    def _build_detect_result(
        self,
        *,
        text: str,
        platform: str,
        mode: str,
        report_summary: dict,
        algo_result,
    ) -> dict:
        base_score = self._heuristic_ai_score(text)
        score, profile, breakdown = self._simulate_platform_detect_score(platform, text, base_score)
        label = "high" if score >= profile["high"] else "medium" if score >= profile["medium"] else "low"

        if isinstance(algo_result, dict):
            result_score = algo_result.get("ai_score")
            if not isinstance(result_score, (float, int)):
                result_score = algo_result.get("aigc_score")
            if isinstance(result_score, (float, int)):
                package_score = float(result_score)
                if package_score > 1:
                    package_score = package_score / 100
                package_score = self._clamp_score(package_score)
                score = round(self._clamp_score(score * 0.65 + package_score * 0.35), 4)
                breakdown["algo_package_score"] = round(package_score, 4)
                breakdown["blended"] = True
            if algo_result.get("label"):
                label = str(algo_result.get("label")).strip().lower()
            elif algo_result.get("level"):
                level_raw = str(algo_result.get("level")).strip().lower()
                level_map = {
                    "高": "high",
                    "中": "medium",
                    "低": "low",
                    "high": "high",
                    "medium": "medium",
                    "low": "low",
                }
                label = level_map.get(level_raw, level_raw)
        else:
            breakdown["blended"] = False

        band = self._risk_band(score, high=profile["high"], medium=profile["medium"])
        risk_paragraphs = self._top_risk_paragraphs(text, platform=platform)
        return {
            "type": TaskType.AIGC_DETECT.value,
            "platform": platform,
            "simulation_profile": profile["name"],
            "mode": mode,
            "llm_used": self._pipeline_usage["llm_used"],
            "algo_package_used": self._pipeline_usage["algo_package_used"],
            "ai_score": score,
            "score_pct": round(score * 100, 2),
            "label": label,
            "risk_band": band,
            "summary": f"AIGC检测完成，当前文本判定为{band}，建议结合高风险段落进行人工复核。",
            "source_stats": self._text_stats(text),
            "report_summary": report_summary,
            "score_breakdown": breakdown,
            "risk_paragraphs": risk_paragraphs,
        }

    def _risk_band(self, score: float, *, high: float = 0.65, medium: float = 0.35) -> str:
        if score >= high:
            return "高风险"
        if score >= medium:
            return "中风险"
        return "低风险"

    def _top_risk_paragraphs(self, text: str, platform: str = "cnki") -> list[dict]:
        paragraphs = [part.strip() for part in text.splitlines() if part.strip()]
        scored = []
        for index, paragraph in enumerate(paragraphs, start=1):
            if len(paragraph) < 20:
                continue
            base_score = self._heuristic_ai_score(paragraph)
            simulated_score, _profile, _breakdown = self._simulate_platform_detect_score(platform, paragraph, base_score)
            scored.append(
                {
                    "index": index,
                    "score": round(simulated_score * 100, 2),
                    "excerpt": self._clip_text(paragraph, 80),
                }
            )
        scored.sort(key=lambda item: item["score"], reverse=True)
        return scored[:3]


    def _write_detect_report_pdf(self, output_path: Path, result: dict) -> None:
        lines = self._build_detect_report_lines(result)
        output_path.write_bytes(self._render_pdf(lines))

    def _build_detect_report_lines(self, result: dict) -> list[str]:
        platform_key = str(result.get("platform") or "cnki").strip().lower()
        platform_label = {
            "cnki": "CNKI-style Simulation",
            "vip": "VIP-style Simulation",
            "paperpass": "PaperPass-style Simulation",
        }.get(platform_key, f"{platform_key.upper()}-style Simulation")
        score_pct = float(result.get("score_pct") or 0.0)
        risk_level = "HIGH" if score_pct >= 65 else "MEDIUM" if score_pct >= 35 else "LOW"
        generated_at = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        stats = result.get("source_stats") or {}

        lines: list[str] = [
            "格物学术 AIGC 检测报告",
            f"Generated At: {generated_at}",
            f"Platform: {platform_label}",
            "Engine Type: Simulated provider algorithm (not official provider API)",
            f"Simulation Profile: {result.get('simulation_profile', platform_key + '_like')}",
            "",
            "1. Detection Summary",
            f"   - Composite Score: {score_pct:.2f}%",
            f"   - Risk Level: {risk_level}",
            f"   - Pipeline Mode: {result.get('mode') or '-'}",
            f"   - Model Assisted: {'YES' if result.get('llm_used') else 'NO'}",
            f"   - Algorithm Package Used: {'YES' if result.get('algo_package_used') else 'NO'}",
            "",
            "2. Source Text Statistics",
            f"   - Characters: {stats.get('char_count', 0)}",
            f"   - Paragraphs: {stats.get('paragraph_count', 0)}",
            f"   - Sentences: {stats.get('sentence_count', 0)}",
            f"   - Avg Sentence Length: {stats.get('avg_sentence_length', 0)}",
            "",
            "3. Supplementary Report Signals",
        ]

        report_summary = result.get("report_summary") or {}
        if report_summary.get("available"):
            metrics = report_summary.get("metrics") or []
            actions = report_summary.get("recommended_actions") or []
            if metrics:
                lines.append("   - Parsed Metrics:")
                for metric in metrics:
                    lines.append(
                        f"     * {metric.get('label', 'Metric')}: {metric.get('value', '-')}{metric.get('unit', '')}"
                    )
            if actions:
                lines.append("   - Suggested Actions:")
                for action in actions[:5]:
                    lines.append(f"     * {action}")
            if not metrics and not actions:
                lines.append("   - Supplementary report uploaded, but no structured metrics were extracted.")
        else:
            lines.append("   - No supplementary report uploaded.")

        lines.extend(["", "4. Top Risk Paragraphs"])
        risk_paragraphs = result.get("risk_paragraphs") or []
        if risk_paragraphs:
            for item in risk_paragraphs:
                lines.append(
                    f"   - Paragraph {item.get('index', '-')}: {item.get('score', 0)}% | {self._clip_text(item.get('excerpt', ''), 120)}"
                )
        else:
            lines.append("   - No high-risk paragraph extracted.")

        lines.extend(
            [
                "",
                "Disclaimer:",
                "This report is generated by an internal simulation engine for operational use.",
                "It is not an official report issued by CNKI, VIP, or PaperPass.",
            ]
        )
        return lines

    def _wrap_pdf_line(self, text: str, width: int = 90) -> list[str]:
        compact = " ".join(str(text or "").split())
        if not compact:
            return [""]
        wrapped: list[str] = []
        remaining = compact
        while len(remaining) > width:
            cut = remaining.rfind(" ", 0, width + 1)
            if cut <= 0:
                cut = width
            wrapped.append(remaining[:cut].strip())
            remaining = remaining[cut:].strip()
        if remaining:
            wrapped.append(remaining)
        return wrapped

    def _pdf_safe_text(self, text: str) -> str:
        return str(text or "").encode("latin-1", "replace").decode("latin-1")

    def _pdf_escape(self, text: str) -> str:
        return text.replace("\\", "\\\\").replace("(", "\\(").replace(")", "\\)")

    def _render_pdf(self, lines: list[str]) -> bytes:
        page_width = 595
        page_height = 842
        margin_x = 46
        margin_top = 60
        margin_bottom = 52
        font_size = 11
        line_height = 15

        expanded_lines: list[str] = []
        for line in lines:
            wrapped = self._wrap_pdf_line(line, width=92)
            expanded_lines.extend(wrapped if wrapped else [""])
        if not expanded_lines:
            expanded_lines = [""]

        usable_height = page_height - margin_top - margin_bottom
        lines_per_page = max(1, int(usable_height // line_height))
        page_line_chunks = [
            expanded_lines[i : i + lines_per_page] for i in range(0, len(expanded_lines), lines_per_page)
        ]

        objects: list[tuple[int, bytes]] = [(1, b"<< /Type /Catalog /Pages 2 0 R >>")]
        page_refs: list[str] = []
        next_obj_id = 3
        font_obj_id = 2 + len(page_line_chunks) * 2 + 1

        for chunk in page_line_chunks:
            page_obj_id = next_obj_id
            content_obj_id = next_obj_id + 1
            next_obj_id += 2
            page_refs.append(f"{page_obj_id} 0 R")

            text_ops: list[str] = []
            for index, line in enumerate(chunk):
                y = page_height - margin_top - index * line_height
                safe_line = self._pdf_escape(self._pdf_safe_text(line))
                text_ops.append(f"BT /F1 {font_size} Tf 1 0 0 1 {margin_x} {y:.2f} Tm ({safe_line}) Tj ET")
            stream_text = "\n".join(text_ops)
            stream_bytes = stream_text.encode("latin-1", "replace")

            page_obj = (
                f"<< /Type /Page /Parent 2 0 R /MediaBox [0 0 {page_width} {page_height}] "
                f"/Resources << /Font << /F1 {font_obj_id} 0 R >> >> /Contents {content_obj_id} 0 R >>"
            ).encode("ascii")
            content_obj = (
                f"<< /Length {len(stream_bytes)} >>\nstream\n".encode("ascii")
                + stream_bytes
                + b"\nendstream"
            )
            objects.append((page_obj_id, page_obj))
            objects.append((content_obj_id, content_obj))

        pages_obj = f"<< /Type /Pages /Count {len(page_line_chunks)} /Kids [{' '.join(page_refs)}] >>".encode("ascii")
        objects.insert(1, (2, pages_obj))
        objects.append((font_obj_id, b"<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica >>"))

        objects.sort(key=lambda item: item[0])
        output = bytearray()
        output.extend(b"%PDF-1.4\n%\xe2\xe3\xcf\xd3\n")
        offsets: dict[int, int] = {}

        for obj_id, obj_body in objects:
            offsets[obj_id] = len(output)
            output.extend(f"{obj_id} 0 obj\n".encode("ascii"))
            output.extend(obj_body)
            output.extend(b"\nendobj\n")

        xref_offset = len(output)
        max_obj_id = max(offsets)
        output.extend(f"xref\n0 {max_obj_id + 1}\n".encode("ascii"))
        output.extend(b"0000000000 65535 f \n")
        for obj_id in range(1, max_obj_id + 1):
            offset = offsets.get(obj_id, 0)
            output.extend(f"{offset:010d} 00000 n \n".encode("ascii"))
        output.extend(f"trailer\n<< /Size {max_obj_id + 1} /Root 1 0 R >>\n".encode("ascii"))
        output.extend(f"startxref\n{xref_offset}\n%%EOF".encode("ascii"))
        return bytes(output)


    def _render_detect_report(self, result: dict) -> str:
        lines = [
            "格物学术 AIGC 检测报告",
            f"平台：{result.get('platform')}",
            f"处理模式：{result.get('mode')}",
            f"综合分值：{result.get('score_pct')}%",
            f"风险等级：{result.get('risk_band')}",
            f"摘要：{result.get('summary')}",
            "",
            "文本概览：",
        ]
        stats = result.get("source_stats", {})
        lines.extend(
            [
                f"- 字符数：{stats.get('char_count', 0)}",
                f"- 段落数：{stats.get('paragraph_count', 0)}",
                f"- 句子数：{stats.get('sentence_count', 0)}",
                f"- 平均句长：{stats.get('avg_sentence_length', 0)}",
            ]
        )
        report_summary = result.get("report_summary") or {}
        if report_summary.get("available"):
            lines.extend(["", "辅助报告信息："])
            for metric in report_summary.get("metrics", []):
                lines.append(f"- {metric.get('label')}：{metric.get('value')}{metric.get('unit', '')}")
            for action in report_summary.get("recommended_actions", []):
                lines.append(f"- {action}")
        risk_paragraphs = result.get("risk_paragraphs") or []
        if risk_paragraphs:
            lines.extend(["", "高风险段落："])
            for item in risk_paragraphs:
                lines.append(f"- 段落{item.get('index')} | 风险 {item.get('score')}% | {item.get('excerpt')}")
        return "\n".join(lines)

    def _build_transform_result(
        self,
        *,
        task_type: TaskType,
        platform: str,
        mode: str,
        source_text: str,
        output_text: str,
        report_summary: dict,
    ) -> dict:
        source_stats = self._text_stats(source_text)
        output_stats = self._text_stats(output_text)
        sample_before = source_text[:4000]
        sample_after = output_text[:4000]
        similarity = SequenceMatcher(None, sample_before, sample_after).ratio() if sample_before or sample_after else 1.0
        change_ratio = round((1 - similarity) * 100, 2)
        task_label = "降重" if task_type == TaskType.DEDUP else "降AIGC率"
        review_points = list(report_summary.get("recommended_actions") or [])
        review_points.append("建议下载结果文档后结合原文进行人工终审。")
        review_points.append("重点检查摘要、结论、数据表述和引用位置。")
        deduped_points = list(dict.fromkeys(review_points))
        return {
            "type": task_type.value,
            "platform": platform,
            "mode": mode,
            "llm_used": self._pipeline_usage["llm_used"],
            "algo_package_used": self._pipeline_usage["algo_package_used"],
            "summary": f"{task_label}任务已完成，本次结果已结合正文与辅助报告生成处理摘要。",
            "source_stats": source_stats,
            "output_stats": output_stats,
            "change_ratio": change_ratio,
            "report_summary": report_summary,
            "review_points": deduped_points[:4],
            "output_preview": self._clip_text(output_text, 220),
        }

    def _clip_text(self, text: str, limit: int) -> str:
        compact = " ".join((text or "").split())
        if len(compact) <= limit:
            return compact
        return f"{compact[:limit].rstrip()}..."
