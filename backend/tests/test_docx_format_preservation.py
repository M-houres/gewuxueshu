import base64
from pathlib import Path
from zipfile import ZipFile
import xml.etree.ElementTree as ET

from docx import Document
from docx.shared import Pt
from sqlalchemy.orm import Session

from app.models import TaskType
from app.services.processing_engine import ProcessingEngine

_VOLATILE_ATTRS = {"paraId", "textId"}


def _local_name(tag: str) -> str:
    return tag.split("}", 1)[-1]


def _xml_shape_signature(node: ET.Element, *, mask_text: bool) -> tuple:
    attrs = []
    for key, value in sorted(node.attrib.items(), key=lambda item: _local_name(item[0])):
        local_key = _local_name(key)
        if local_key.startswith("rsid") or local_key in _VOLATILE_ATTRS:
            continue
        attrs.append((local_key, value))
    text = (node.text or "").strip()
    if mask_text and _local_name(node.tag) == "t":
        text = "__TEXT__"
    children = tuple(_xml_shape_signature(child, mask_text=mask_text) for child in list(node))
    return (_local_name(node.tag), tuple(attrs), text, children)


def _word_xml_parts(path: Path) -> dict[str, bytes]:
    with ZipFile(path, "r") as archive:
        return {
            name: archive.read(name)
            for name in archive.namelist()
            if name.startswith("word/") and name.endswith(".xml")
        }


def _content_parts(parts: dict[str, bytes]) -> set[str]:
    result = set()
    for name in parts:
        if name == "word/document.xml" or name.startswith("word/header") or name.startswith("word/footer"):
            result.add(name)
    if "word/footnotes.xml" in parts:
        result.add("word/footnotes.xml")
    return result


def _create_complex_docx(path: Path, image_path: Path) -> None:
    doc = Document()

    section = doc.sections[0]
    section.header.paragraphs[0].text = "页眉: 首先进行格式验证"
    section.footer.paragraphs[0].text = "页脚: 因此需要比对结构"

    p = doc.add_paragraph()
    run_1 = p.add_run("首先，我们构造复杂文档。")
    run_1.bold = True
    run_1.font.name = "宋体"
    run_1.font.size = Pt(14)

    run_2 = p.add_run("因此这段文字会被替换。")
    run_2.italic = True
    run_2.font.name = "Times New Roman"
    run_2.font.size = Pt(12)

    table = doc.add_table(rows=2, cols=2)
    table.style = "Table Grid"
    table.cell(0, 0).text = "其次，表格内容A"
    table.cell(0, 1).text = "总之，表格内容B"
    table.cell(1, 0).text = "研究表明，表格内容C"
    table.cell(1, 1).text = "可以看出，表格内容D"

    doc.add_paragraph("非常重要：保留段落、字体、样式和图片。")
    doc.add_picture(str(image_path))

    doc.save(str(path))


def _collect_body_run_styles(doc: Document) -> list[tuple]:
    styles = []
    for para in doc.paragraphs:
        for run in para.runs:
            styles.append(
                (
                    run.bold,
                    run.italic,
                    run.underline,
                    run.font.name,
                    run.font.size.pt if run.font.size else None,
                )
            )
    return styles


def _collect_table_text(doc: Document) -> list[str]:
    values: list[str] = []
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                values.append(cell.text)
    return values


def _collect_header_footer_text(doc: Document) -> tuple[list[str], list[str]]:
    headers: list[str] = []
    footers: list[str] = []
    for section in doc.sections:
        headers.extend([p.text for p in section.header.paragraphs])
        footers.extend([p.text for p in section.footer.paragraphs])
    return headers, footers


def test_docx_transform_only_changes_text_nodes(db_session: Session, tmp_path: Path) -> None:
    image_bytes = base64.b64decode("iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAQAAAC1HAwCAAAAC0lEQVR42mP8/x8AAwMCAO+nm2QAAAAASUVORK5CYII=")
    image_path = tmp_path / "tiny.png"
    image_path.write_bytes(image_bytes)

    source_path = tmp_path / "source.docx"
    output_path = tmp_path / "output.docx"
    _create_complex_docx(source_path, image_path)

    engine = ProcessingEngine(db_session)
    engine._transform_docx(source_path, output_path, TaskType.DEDUP, "cnki")

    before_doc = Document(str(source_path))
    after_doc = Document(str(output_path))

    assert len(before_doc.tables) == len(after_doc.tables)
    assert len(before_doc.inline_shapes) == len(after_doc.inline_shapes)
    assert _collect_body_run_styles(before_doc) == _collect_body_run_styles(after_doc)
    assert "".join(p.text for p in before_doc.paragraphs) != "".join(p.text for p in after_doc.paragraphs)

    assert _collect_table_text(before_doc) != _collect_table_text(after_doc)
    before_headers, before_footers = _collect_header_footer_text(before_doc)
    after_headers, after_footers = _collect_header_footer_text(after_doc)
    assert before_headers == after_headers
    assert before_footers == after_footers

    before_parts = _word_xml_parts(source_path)
    after_parts = _word_xml_parts(output_path)
    assert set(before_parts.keys()) == set(after_parts.keys())

    text_parts = _content_parts(before_parts)
    for name in before_parts:
        before_root = ET.fromstring(before_parts[name])
        after_root = ET.fromstring(after_parts[name])
        mask_text = name in text_parts
        assert _xml_shape_signature(before_root, mask_text=mask_text) == _xml_shape_signature(after_root, mask_text=mask_text)

    assert before_parts["word/document.xml"] != after_parts["word/document.xml"]
