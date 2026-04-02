from pathlib import Path

from docx import Document

from app.utils import extract_text_from_file


def test_extract_text_from_docx_includes_table_cells(tmp_path: Path) -> None:
    doc_path = tmp_path / "with_table.docx"
    doc = Document()
    doc.add_paragraph("正文段落")
    table = doc.add_table(rows=1, cols=2)
    table.cell(0, 0).text = "表格字段A"
    table.cell(0, 1).text = "表格字段B"
    doc.save(str(doc_path))

    text = extract_text_from_file(doc_path)
    assert "正文段落" in text
    assert "表格字段A" in text
    assert "表格字段B" in text
